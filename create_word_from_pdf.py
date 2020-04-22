from PIL import Image
from pdf2image import convert_from_path
import glob 
from pathlib import Path
import shutil, os
from docx import Document
import fnmatch
import re
import shutil


def find_files_ignore_case(which, where='.'):
    '''Returns list of filenames from `where` path matched by 'which'
       shell pattern. Matching is case-insensitive.'''
    
    # TODO: recursive param with walk() filtering
    rule = re.compile(fnmatch.translate(which), re.IGNORECASE)
    return [name for name in os.listdir(where) if rule.match(name)]


def crop_image_center(file, crop_left, crop_right, crop_top, crop_bottom):
    img = Image.open(file)
    x, y = img.size
    box = (crop_left, crop_top, x - crop_left - crop_right, y - crop_top - crop_bottom)
    crop = img.crop(box)    
    crop.save(file)

def create_empty_folder(path):
    '''Create a folder. Delete content if exists'''
    Path(path).mkdir(parents=True, exist_ok=True)
    
    # Remove existing files
    existing_files = find_files_ignore_case(os.path.join(path, '*'))
    for ef in existing_files:
        os.remove(ef)

def convert_pdf_to_images(file):
    '''Convert a PDF file into images and save to folder of same name
        Return folder which contains the images
    '''
    # Create directory for each file
    folder = os.path.splitext(file)[0]    
    create_empty_folder(folder)
    
    # Convert PDF to images into the directory
    images = convert_from_path(file)
    for i, image in enumerate(images):
        file_name = 'Z{:05}.jpg'.format(i+1)
        image.save(os.path.join(folder, file_name), 'JPEG')

    return folder

def get_file_name_prefix(filename):
    with open('file_name_prefixes.txt') as f:
        for line in f:
            line = line.strip()
            if filename.lower().startswith(line.lower()):
                return line.strip()
    return None

import errno, os, stat, shutil
def handleRemoveReadonly(func, path, exc):
  excvalue = exc[1]
  if func in (os.rmdir, os.remove) and excvalue.errno == errno.EACCES:
      os.chmod(path, stat.S_IRWXU| stat.S_IRWXG| stat.S_IRWXO) # 0777
      func(path)
  else:
      raise
      

if __name__ == '__main__':
    cur_folder = os.path.abspath('')
    
    # Convert PDFs to Images
    print('Convert PDFs to images...')
    files = find_files_ignore_case('*.pdf')
    for pdf_file in files:
        pdf_file = os.path.join(cur_folder, pdf_file)
        print(pdf_file)
        folder = convert_pdf_to_images(pdf_file)
    
    # Crop images
    print('Crop images...')
    files = find_files_ignore_case('*.pdf')
    for file in files:
        folder = os.path.splitext(file)[0]
        print(folder)
        images = find_files_ignore_case('*.jpg', folder)
        images.sort()
        print(images)
        for image_file in images:
            try:
                image_file = os.path.join(folder, image_file)
                crop_image_center(image_file, crop_left=160, 
                        crop_right=-40, crop_top=100, crop_bottom=20)
            except:
                pass
    
    # Copy Image *.jpg From Reference to Folder
    files = find_files_ignore_case('*.pdf')
    for file in files:
        print(file)
        folder = os.path.splitext(file)[0]
        file_prefix = get_file_name_prefix(file)
        print(file_prefix)

        # Copy Image *.jpg From Reference to Folder
        source_files = find_files_ignore_case('{}*.jpg'.format(file_prefix), 'Reference')

        for f in source_files:
            f = os.path.join('Reference', f)
            shutil.copy(f, folder)
    
    # Insert Images to Word
    files = find_files_ignore_case('*.pdf')
    for file in files:
        folder = os.path.splitext(file)[0] 
        word_file = folder+".docx"

        # Copy from template docx
        file_prefix = get_file_name_prefix(file)
        files = find_files_ignore_case('{}*.docx'.format(file_prefix), 'Reference')
        print(file, file_prefix, files)
        if files:
            document = Document(os.path.join('Reference', files[0]))
            document.add_section()
        else:
            document = Document()
        document.save(word_file)

        section = document.sections[0]
    #         width = section.page_width - section.left_margin - section.right_margin
        height = section.page_height - section.top_margin - section.bottom_margin

        images = find_files_ignore_case('*.jpg', folder)
        for image_file in images:
            image_file = os.path.join(folder, image_file)
    #         document.add_picture(image_file, width=width)
            document.add_picture(image_file, height=height)

        document.save(word_file)
            
    # Delete folders including its images
    files = find_files_ignore_case('*.pdf')
    for file in files:
        folder = os.path.splitext(file)[0]
        print('Deleting', folder, os.path.isdir(folder))
        try:
            files_in_dir = os.listdir(folder) 
            for file in files_in_dir:                  # loop to delete each file in folder
                os.remove(os.path.join(folder,file))
            #os.rmdir(folder)
            shutil.rmtree(folder, ignore_errors=False, onerror=handleRemoveReadonly)
        except Exception as ex:
            print('Error deleting', folder, ex)

